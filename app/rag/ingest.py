import os
import uuid
from datetime import date
from pypdf import PdfReader
import docx

from sqlalchemy.orm import Session

from app.models import Document, Chunk
from app.utils.text_splitter import chunk_text, ChunkConfig
from app.utils.hashing import sha256_text

def _read_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)

def _read_docx(path: str) -> str:
    d = docx.Document(path)
    return "\n".join(p.text for p in d.paragraphs)

def _read_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def load_file_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _read_pdf(path)
    if ext == ".docx":
        return _read_docx(path)
    if ext in [".txt", ".md", ".csv", ".log"]:
        return _read_text(path)
    # skip unknown
    return ""

def infer_doc_type(path: str) -> str:
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    return ext or "unknown"

def ingest_folder(
    db: Session,
    folder_path: str,
    embed_fn,
    product_line: str | None = None,
    region: str | None = None,
    doc_type: str | None = None,
    effective_date: str | None = None,
    chunk_cfg: ChunkConfig | None = None,
) -> tuple[int, int, int]:
    """
    Returns: (docs_added, docs_skipped, chunks_added)
    """
    chunk_cfg = chunk_cfg or ChunkConfig()

    docs_added = 0
    docs_skipped = 0
    chunks_added = 0

    eff_dt = None
    if effective_date:
        y, m, d = map(int, effective_date.split("-"))
        eff_dt = date(y, m, d)

    for root, _, files in os.walk(folder_path):
        for name in files:
            path = os.path.join(root, name)
            text = load_file_text(path).strip()
            if not text:
                continue

            h = sha256_text(text)
            exists = db.query(Document).filter(Document.content_hash == h).first()
            if exists:
                docs_skipped += 1
                continue

            doc = Document(
                id=uuid.uuid4(),
                title=os.path.basename(path),
                source_path=os.path.abspath(path),
                doc_type=doc_type or infer_doc_type(path),
                product_line=product_line,
                region=region,
                effective_date=eff_dt,
                content_hash=h,
            )
            db.add(doc)
            db.flush()  # get doc.id

            chunks = chunk_text(text, chunk_cfg)
            if not chunks:
                docs_skipped += 1
                db.rollback()
                continue

            # embed in batches for efficiency
            embeddings = embed_fn(chunks)

            for i, (ch, emb) in enumerate(zip(chunks, embeddings)):
                c = Chunk(
                    id=uuid.uuid4(),
                    document_id=doc.id,
                    chunk_index=i,
                    text=ch,
                    embedding=emb,
                    char_count=len(ch),
                    metadata_={
                        "source_path": doc.source_path,
                        "title": doc.title,
                        "doc_type": doc.doc_type,
                        "product_line": doc.product_line,
                        "region": doc.region,
                        "effective_date": str(doc.effective_date) if doc.effective_date else None,
                    },
                )
                db.add(c)
                chunks_added += 1

            docs_added += 1

    db.commit()
    return docs_added, docs_skipped, chunks_added